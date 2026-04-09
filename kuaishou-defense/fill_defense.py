#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
快手答辩文档自动填写工具
支持【】占位符格式，自动应用宋体小四字体
"""

import os
import sys
import argparse
from datetime import datetime
from pathlib import Path
import subprocess
import tempfile
import shutil

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("❌ 缺少依赖: python-docx")
    print("请运行: pip3 install python-docx")
    sys.exit(1)


def convert_doc_to_docx(doc_path):
    """将.doc文件转换为.docx (仅macOS)"""
    doc_path = Path(doc_path)
    temp_dir = Path(tempfile.mkdtemp())
    temp_docx = temp_dir / f"{doc_path.stem}.docx"
    
    try:
        result = subprocess.run(
            ["textutil", "-convert", "docx", "-output", str(temp_docx), str(doc_path)],
            capture_output=True,
            text=True
        )
        
        if result.returncode == 0 and temp_docx.exists():
            return temp_docx
        else:
            print(f"⚠️  转换失败: {doc_path.name}")
            return None
    except Exception as e:
        print(f"⚠️  转换出错: {e}")
        return None


def replace_text_in_doc(doc, replacements):
    """在Word文档中替换文本（整段替换，保持原格式）"""
    count = 0
    
    # 替换段落中的文本
    for para in doc.paragraphs:
        original_text = para.text
        new_text = original_text
        
        # 删除底部页码文本 (PAGE1 / NUMPAGES1)
        if 'PAGE' in original_text and 'NUMPAGES' in original_text:
            # 清空整个段落
            for run in para.runs:
                run.text = ''
            continue
        
        # 执行所有替换
        for old_text, new_text_replacement in replacements.items():
            if old_text in new_text:
                new_text = new_text.replace(old_text, new_text_replacement)
                count += 1
        
        # 如果有替换，更新整个段落（不修改字体）
        if new_text != original_text:
            for run in para.runs:
                run.text = ''
            
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
    
    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original_text = para.text
                    new_text = original_text
                    
                    # 删除表格中的页码文本
                    if 'PAGE' in original_text and 'NUMPAGES' in original_text:
                        for run in para.runs:
                            run.text = ''
                        continue
                    
                    for old_text, new_text_replacement in replacements.items():
                        if old_text in new_text:
                            new_text = new_text.replace(old_text, new_text_replacement)
                            count += 1
                    
                    if new_text != original_text:
                        for run in para.runs:
                            run.text = ''
                        
                        if para.runs:
                            para.runs[0].text = new_text
                        else:
                            para.add_run(new_text)
    
    return count


class KuaishouDefenseFiller:
    """快手答辩文档填写器"""
    
    def __init__(self, templates_dir=None, output_dir=None):
        self.script_dir = Path(__file__).parent
        self.templates_dir = Path(templates_dir) if templates_dir else self.script_dir / "templates"
        self.output_dir = Path(output_dir) if output_dir else Path.home() / "Downloads"
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def fill_all(self, plaintiff, defendant, case_number, shop_name,
                 business_entity, order_number, court, date=None):
        """填写所有文档"""
        
        date_str = date if date else datetime.now().strftime("%Y年%m月%d日")
        
        print(f"\n📝 开始填写快手答辩文档...")
        print(f"原告: {plaintiff}")
        print(f"被告: {defendant}")
        print(f"案号: {case_number}")
        print(f"日期: {date_str}")
        print("-" * 60)
        
        results = []
        
        # 1. 授权委托书
        auth_template = self.templates_dir / "1-快手-授权委托书.docx"
        if auth_template.exists():
            doc = Document(auth_template)
            replacements = {
                "【原告】": plaintiff,
                "【被告】": defendant,
                "【案号】": case_number,
                "【年月日】": date_str
            }
            count = replace_text_in_doc(doc, replacements)
            output_file = self.output_dir / f"授权委托书_{case_number.replace('/', '-')}_{datetime.now().strftime('%Y%m%d')}.docx"
            doc.save(output_file)
            print(f"✅ 授权委托书已生成 (替换{count}处): {output_file.name}")
            results.append(output_file)
        
        # 2. 任职证明
        emp_template = self.templates_dir / "2-快手-任职证明.docx"
        if emp_template.exists():
            doc = Document(emp_template)
            replacements = {"【年月日】": date_str}
            count = replace_text_in_doc(doc, replacements)
            output_file = self.output_dir / f"任职证明_{datetime.now().strftime('%Y%m%d')}.docx"
            doc.save(output_file)
            print(f"✅ 任职证明已生成 (替换{count}处): {output_file.name}")
            results.append(output_file)
        
        # 3. 答辩状
        def_template = self.templates_dir / "3-快手-答辩状.docx"
        if def_template.exists():
            doc = Document(def_template)
            
            # 根据订单号是否填写,决定替换方式
            if order_number:
                # 有订单号:正常替换
                replacements = {
                    "【原告】": plaintiff,
                    "【被告】": defendant,
                    "【案号】": case_number,
                    "【店铺名称】": shop_name,
                    "【经营主体】": business_entity,
                    "【订单号】": order_number,
                    "【XXX法院】": court,
                    "【年月日】": date_str,
                    "年月日": date_str  # 支持无括号的占位符
                }
            else:
                # 无订单号:删除"订单号【订单号】"整个字符串
                replacements = {
                    "【原告】": plaintiff,
                    "【被告】": defendant,
                    "【案号】": case_number,
                    "【店铺名称】": shop_name,
                    "【经营主体】": business_entity,
                    "订单号【订单号】": "",  # 删除"订单号【订单号】"
                    "【XXX法院】": court,
                    "【年月日】": date_str,
                    "年月日": date_str  # 支持无括号的占位符
                }
            
            count = replace_text_in_doc(doc, replacements)
            output_file = self.output_dir / f"答辩状_{case_number.replace('/', '-')}_{datetime.now().strftime('%Y%m%d')}.docx"
            doc.save(output_file)
            print(f"✅ 答辩状已生成 (替换{count}处): {output_file.name}")
            results.append(output_file)
        
        print("-" * 60)
        print(f"✨ 完成! 共生成 {len(results)} 个文档")
        print(f"输出目录: {self.output_dir}")
        
        return results


def main():
    parser = argparse.ArgumentParser(description="快手答辩文档自动填写工具")
    parser.add_argument("--plaintiff", required=True, help="原告名称")
    parser.add_argument("--defendant", required=True, help="被告名称")
    parser.add_argument("--case-number", required=True, help="案号")
    parser.add_argument("--shop-name", required=True, help="店铺名称")
    parser.add_argument("--business-entity", required=True, help="经营主体")
    parser.add_argument("--order-number", required=True, help="订单编号")
    parser.add_argument("--court", required=True, help="法院名称")
    parser.add_argument("--date", help="日期 (格式: YYYY年MM月DD日)")
    parser.add_argument("--output-dir", help="输出目录")
    
    args = parser.parse_args()
    
    filler = KuaishouDefenseFiller(output_dir=args.output_dir)
    results = filler.fill_all(
        plaintiff=args.plaintiff,
        defendant=args.defendant,
        case_number=args.case_number,
        shop_name=args.shop_name,
        business_entity=args.business_entity,
        order_number=args.order_number,
        court=args.court,
        date=args.date
    )
    
    return 0 if results else 1


if __name__ == "__main__":
    sys.exit(main())
